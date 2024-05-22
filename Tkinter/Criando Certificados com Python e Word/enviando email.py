import pandas as pd
from tkinter import *
from tkinter import messagebox
from tkinter import ttk

from docx import Document
from docx.shared import Pt

import win32com.client as win32

import pandas as pd
from datetime import date
import numpy as np


outlook = win32.Dispatch("outlook.application")


janela = Tk()

janela.title("Gerador de Certificados")

stilo = ttk.Style()
stilo.theme_use("alt")
stilo.configure(".", font=("Arial 20"), rowheight=30)

treeviewdados = ttk.Treeview(janela, columns=(1,2,3,4,5,6), show="headings")

treeviewdados.column("1", anchor=CENTER)
treeviewdados.heading("1", text="CPF")

treeviewdados.column("2", anchor=CENTER)
treeviewdados.heading("2", text="Nome")

treeviewdados.column("3", anchor=CENTER)
treeviewdados.heading("3", text="RG")

treeviewdados.column("4", anchor=CENTER)
treeviewdados.heading("4", text="Data de início")

treeviewdados.column("5", anchor=CENTER)
treeviewdados.heading("5", text="Data de conclusão")

treeviewdados.column("6", anchor=CENTER)
treeviewdados.heading("6", text="E-mail")

treeviewdados.grid(row=4, column=0, columnspan=6, sticky="NSEW", pady=15)

def funcaopassadadostreevireparaentry(Event):

    item = treeviewdados.selection()

    for i in item:

        exibircpf.delete(0, END)
        exibirnome.delete(0, END)
        exibirrg.delete(0, END)
        exibirinicio.delete(0, END)
        exibirfim.delete(0, END)
        exibiremail.delete(0, END)

        exibircpf.insert(0, treeviewdados.item(i, "values")[0])
        exibirnome.insert(0, treeviewdados.item(i, "values")[1])
        exibirrg.insert(0, treeviewdados.item(i, "values")[2])
        exibirinicio.insert(0, treeviewdados.item(i, "values")[3])
        exibirfim.insert(0, treeviewdados.item(i, "values")[4])
        exibiremail.insert(0, treeviewdados.item(i, "values")[5])

treeviewdados.bind("<Double-1>", funcaopassadadostreevireparaentry)

dadousuarios = pd.read_excel("C:\\Users\\georg\\Documents\\Python - GitHub\\Python\\Tkinter\\Criando Certificados com Python e Word\\Dados.xlsx")

dadousuarios["Data Inicio"] = dadousuarios["Data Inicio"].astype(str)
dadousuarios["Data Fim"] = dadousuarios["Data Fim"].astype(str)


for linha in range(len(dadousuarios)):

    datainiciopt1 = dadousuarios.iloc[linha, 3].split("-")[0]
    datainiciopt2 = dadousuarios.iloc[linha, 3].split("-")[1]
    datainiciopt3 = dadousuarios.iloc[linha, 3].split("-")[2]

    datainiciotratada = datainiciopt3 + "/" + datainiciopt2 + "/" + datainiciopt1

    print(datainiciotratada)

    datafimpt1 = dadousuarios.iloc[linha, 4].split("-")[0]
    datafimpt2 = dadousuarios.iloc[linha, 4].split("-")[1]
    datafimpt3 = dadousuarios.iloc[linha, 4].split("-")[2]

    datafimtratada = datafimpt1 + "/" + datafimpt2 + "/" + datafimpt3

    print(datafimtratada)

    treeviewdados.insert("", "end",
                         values=(str(dadousuarios.iloc[linha, 0]),
                                 str(dadousuarios.iloc[linha, 1]),
                                 str(dadousuarios.iloc[linha, 2]),
                                 str(datainiciotratada),
                                 str(datafimtratada),
                                 str(dadousuarios.iloc[linha, 5])))

cpf = Label(text="CPF:", font="Arial 12")
cpf.grid(row=0, column=0, sticky="E", pady=15)

exibircpf = Entry(font="Arial 12")
exibircpf.grid(row=0, column=1, sticky="W", pady=15)

nome = Label(text="Nome:", font="Arial 12")
nome.grid(row=0, column=2, sticky="E", pady=15)

exibirnome = Entry(font="Arial 12")
exibirnome.grid(row=0, column=3, sticky="W", pady=15)

rg = Label(text="RG:", font="Arial 12")
rg.grid(row=0, column=4, sticky="E", pady=15)

exibirrg = Entry(font="Arial 12")
exibirrg.grid(row=0, column=5, sticky="W", pady=15)

datainicio = Label(text="Data de início:", font="Arial 12")
datainicio.grid(row=1, column=0, sticky="E", pady=15)

exibirinicio = Entry(font="Arial 12")
exibirinicio.grid(row=1, column=1, sticky="W", pady=15)

datafim = Label(text="Data de conclusão:", font="Arial 12")
datafim.grid(row=1, column=2, sticky="E", pady=15)

exibirfim = Entry(font="Arial 12")
exibirfim.grid(row=1, column=3, sticky="W", pady=15)

email = Label(text="E-mail:", font="Arial 12")
email.grid(row=1, column=4, sticky="E", pady=15)

exibiremail = Entry(font="Arial 12")
exibiremail.grid(row=1, column=5, sticky="W", pady=15)

def filtrardados():

    for linha in range(len(dadousuarios)):

        todaslinhas = treeviewdados.get_children()

        treeviewdados.delete(*todaslinhas)

        if exibircpf.get() == "":

            botaopesquisar.config(text="Filtrar")

            for linha in range(len(dadousuarios)):

                datainiciopt1 = dadousuarios.iloc[linha, 3].split("-")[0]
                datainiciopt2 = dadousuarios.iloc[linha, 3].split("-")[1]
                datainiciopt3 = dadousuarios.iloc[linha, 3].split("-")[2]

                datainiciotratada = datainiciopt3 + "/" + datainiciopt2 + "/" + datainiciopt1

                print(datainiciotratada)

                datafimpt1 = dadousuarios.iloc[linha, 4].split("-")[0]
                datafimpt2 = dadousuarios.iloc[linha, 4].split("-")[1]
                datafimpt3 = dadousuarios.iloc[linha, 4].split("-")[2]

                datafimtratada = datafimpt1 + "/" + datafimpt2 + "/" + datafimpt3

                print(datafimtratada)

                treeviewdados.insert("", "end",
                                     values=(str(dadousuarios.iloc[linha, 0]),
                                             str(dadousuarios.iloc[linha, 1]),
                                             str(dadousuarios.iloc[linha, 2]),
                                             str(datainiciotratada),
                                             str(datafimtratada),
                                             str(dadousuarios.iloc[linha, 5])))
        else:

            botaopesquisar.config(text="Limpar Filtros")

            for linha in range(len(dadousuarios)):

                datainiciopt1 = dadousuarios.iloc[linha, 3].split("-")[0]
                datainiciopt2 = dadousuarios.iloc[linha, 3].split("-")[1]
                datainiciopt3 = dadousuarios.iloc[linha, 3].split("-")[2]

                datainiciotratada = datainiciopt3 + "/" + datainiciopt2 + "/" + datainiciopt1

                print(datainiciotratada)

                datafimpt1 = dadousuarios.iloc[linha, 4].split("-")[0]
                datafimpt2 = dadousuarios.iloc[linha, 4].split("-")[1]
                datafimpt3 = dadousuarios.iloc[linha, 4].split("-")[2]

                datafimtratada = datafimpt1 + "/" + datafimpt2 + "/" + datafimpt3

                print(datafimtratada)

                if exibircpf.get() == str(dadousuarios.iloc[linha, 0]):

                    treeviewdados.insert("", "end",
                                         values=(str(dadousuarios.iloc[linha, 0]),
                                                 str(dadousuarios.iloc[linha, 1]),
                                                 str(dadousuarios.iloc[linha, 2]),
                                                 str(datainiciotratada),
                                                 str(datafimtratada),
                                                 str(dadousuarios.iloc[linha, 5])))


botaopesquisar = Button(text="PESQUISAR",font="Arial 14",command=filtrardados)
botaopesquisar.grid(row=5, column=0, columnspan=2, sticky="NSEW", padx=20)

def gerarcertificado():

    arquivoword = Document("C:\\Users\\georg\\Documents\\Python - GitHub\\Python\\Tkinter\\Criando Certificados com Python e Word\\Certificado.docx")

    estilo = arquivoword.styles["Normal"]

    nomealuno = exibirnome.get()
    dataInicio = exibirinicio.get()
    datafinal = exibirfim.get()
    nomeinstrutor = "Clevilson Santos"
    cpfaluno = exibircpf.get()
    rgaluno = exibirrg.get()

    frase_pt1 = " concluiu com sucesso o curso de Python RPA, com a carga horária de 20 horas, promovido pela escola de Cursos Online de "

    frasemontada = f"{nomealuno}, CPF: {cpfaluno}, RG: {rgaluno}, {frase_pt1} {dataInicio} a {datafinal}."

    for paragrafo in arquivoword.paragraphs:

        if "@nome" in paragrafo.text:

            paragrafo.text = nomealuno
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

        if "@DataFim" in paragrafo.text:
            paragrafo.text = frasemontada
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

    caminhocertificadogerado = "C:\\Users\\georg\\PycharmProjects\\Tkinter\\"+ nomealuno+ ".docx"

    arquivoword.save(caminhocertificadogerado)

    exibircpf.delete(0, END)
    exibirnome.delete(0, END)
    exibirrg.delete(0, END)
    exibirinicio.delete(0, END)
    exibirfim.delete(0, END)
    exibiremail.delete(0, END)

    messagebox.showinfo("Mensagem", "Certificado gerado com sucesso!")

botaogerarcertificado = Button(text="GERAR CERTIFICADO",font="Arial 14",command=gerarcertificado)
botaogerarcertificado.grid(row=5, column=2, columnspan=2, sticky="NSEW", padx=20)

def gerarcertificadoemmassa():

    for linha in treeviewdados.get_children():

        coluna = treeviewdados.item(linha)["values"]

        cpfseparado = coluna[0]
        nomeAluno = coluna[1]
        rgreparado = coluna[2]
        datainicioseparado = coluna[3]
        datafimseparado = coluna[4]
        nomeinstrutirseparado = "Clevison Santos"

        arquivoword = Document("C:\\Users\\georg\\Downloads\\Certificado.docx")

        estilo = arquivoword.styles["Normal"]

        frase_pt1 = " concluiu com sucesso o curso de Python RPA, com a carga horária de 20 horas, promovido pela escola de Cursos Online de "

        frasemontada = f"{nomeAluno}, CPF: {cpfseparado}, RG: {rgreparado}, {frase_pt1} {datainicioseparado} a {datafimseparado}."

        for paragrafo in arquivoword.paragraphs:

            if "@nome" in paragrafo.text:
                paragrafo.text = nomeAluno
                fonte = estilo.font
                fonte.name = "Calibri (Corpo)"
                fonte.size = Pt(24)

            if "@DataFim" in paragrafo.text:
                paragrafo.text = frasemontada
                fonte = estilo.font
                fonte.name = "Calibri (Corpo)"
                fonte.size = Pt(24)

        caminhocertificadogerado = "C:\\Users\\georg\\PycharmProjects\\Tkinter\\" + nomeAluno + ".docx"

        arquivoword.save(caminhocertificadogerado)

        exibircpf.delete(0, END)
        exibirnome.delete(0, END)
        exibirrg.delete(0, END)
        exibirinicio.delete(0, END)
        exibirfim.delete(0, END)
        exibiremail.delete(0, END)

    messagebox.showinfo("Mensagem", "Certificados gerados com sucesso!")


botaogerarcertificadoemmassa = Button(text="GERAR EM MASSA",font="Arial 14",command=gerarcertificadoemmassa)
botaogerarcertificadoemmassa.grid(row=5, column=4, columnspan=2, sticky="NSEW", padx=20)


def criaremail():

    arquivoword = Document("C:\\Users\\georg\\Downloads\\Certificado.docx")

    estilo = arquivoword.styles["Normal"]

    nomealuno = exibirnome.get()
    dataInicio = exibirinicio.get()
    datafinal = exibirfim.get()
    nomeinstrutor = "Clevilson Santos"
    cpfaluno = exibircpf.get()
    rgaluno = exibirrg.get()
    emailaluno = exibiremail.get()

    frase_pt1 = " concluiu com sucesso o curso de Python RPA, com a carga horária de 20 horas, promovido pela escola de Cursos Online de "

    frasemontada = f"{nomealuno}, CPF: {cpfaluno}, RG: {rgaluno}, {frase_pt1} {dataInicio} a {datafinal}."

    for paragrafo in arquivoword.paragraphs:

        if "@nome" in paragrafo.text:
            paragrafo.text = nomealuno
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

        if "@DataFim" in paragrafo.text:
            paragrafo.text = frasemontada
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

    caminhocertificadogerado = "C:\\Users\\georg\\PycharmProjects\\Tkinter\\" + nomealuno + ".docx"

    arquivoword.save(caminhocertificadogerado)

    exibircpf.delete(0, END)
    exibirnome.delete(0, END)
    exibirrg.delete(0, END)
    exibirinicio.delete(0, END)
    exibirfim.delete(0, END)

    emailoutlook = outlook.CreateItem(0)

    variavelnome = nomealuno.split(" ")[0]

    emailoutlook.To = emailaluno
    emailoutlook.Subject = "Certificado de conclusão de curso " + str(nomealuno)
    emailoutlook.HTMLBody = f"""
    <p>Parabéns, <b>{variavelnome}</b>!</p>
    <p>Você concluiu o curso!</p>
    <p>Segue em anexo o seu certificado de conclusão.</p>
    <p>Atenciosamente.</p>
    """

    emailoutlook.Attachments.Add(caminhocertificadogerado)

    emailoutlook.save()

    exibircpf.delete(0, END)
    exibirnome.delete(0, END)
    exibirrg.delete(0, END)
    exibirinicio.delete(0, END)
    exibirfim.delete(0, END)
    exibiremail.delete(0, END)

    messagebox.showinfo(title="Atenção!", message="Email criado com sucesso!")

botaoemail =Button(text="Criar Email",
                     font="Arial 20",
                     command=criaremail)

botaoemail.grid(row=1, column=6, columnspan=2, sticky="NSEW")

def criaremailmassa():

    for linha in treeviewdados.get_children():
        coluna = treeviewdados.item(linha)["values"]
        cpfseparado = coluna[0]
        nomeAluno = coluna[1]
        rgreparado = coluna[2]
        datainicioseparado = coluna[3]
        datafimseparado = coluna[4]
        nomeinstrutirseparado = "Clevison Santos"
        emailaluno = coluna[5]

        arquivoword = Document("C:\\Users\\georg\\Downloads\\Certificado.docx")
        estilo = arquivoword.styles["Normal"]

        frase_pt1 = " concluiu com sucesso o curso de Python RPA, com a carga horária de 20 horas, promovido pela escola de Cursos Online de "
        frasemontada = f"{nomeAluno}, CPF: {cpfseparado}, RG: {rgreparado}, {frase_pt1} {datainicioseparado} a {datafimseparado}."

        for paragrafo in arquivoword.paragraphs:
            if "@nome" in paragrafo.text:
                paragrafo.text = nomeAluno
                fonte = estilo.font
                fonte.name = "Calibri (Corpo)"
                fonte.size = Pt(24)

            if "@DataFim" in paragrafo.text:
                paragrafo.text = frasemontada
                fonte = estilo.font
                fonte.name = "Calibri (Corpo)"
                fonte.size = Pt(24)

        caminhocertificadogerado = "C:\\Users\\georg\\PycharmProjects\\Tkinter\\" + nomeAluno + ".docx"
        arquivoword.save(caminhocertificadogerado)

        emailoutlook = outlook.CreateItem(0)

        variavelnome = nomeAluno.split(" ")[0]

        emailoutlook.To = emailaluno
        emailoutlook.Subject = "Certificado de conclusão de curso " + str(nomeAluno)
        emailoutlook.HTMLBody = f"""
                 <p>Parabéns, <b>{variavelnome}</b>!</p>
                 <p>Você concluiu o curso!</p>
                 <p>Segue em anexo o seu certificado de conclusão.</p>
                 <p>Atenciosamente.</p>
                 """

        emailoutlook.Attachments.Add(caminhocertificadogerado)

        emailoutlook.save()

            # Limpar os campos após enviar os e-mails
        exibircpf.delete(0, END)
        exibirnome.delete(0, END)
        exibirrg.delete(0, END)
        exibirinicio.delete(0, END)
        exibirfim.delete(0, END)
        exibiremail.delete(0, END)

    messagebox.showinfo(title="Atenção!", message="Emails criados com sucesso!")


botaoemailmassa =Button(text="Criar Emails em Massa",
                     font="Arial 20",
                     command=criaremailmassa)

botaoemailmassa.grid(row=2, column=6, columnspan=2, sticky="NSEW")


janela.mainloop()

