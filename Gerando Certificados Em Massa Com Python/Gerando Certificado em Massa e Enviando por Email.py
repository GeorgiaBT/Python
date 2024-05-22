from docx import Document
from docx.shared import Pt

from docx.shared import RGBColor

from openpyxl import load_workbook

import win32com.client as win32
outlook = win32.Dispatch("outlook.application")


arquivoalunos = "C:\\Users\\georg\\Documents\\Python - GitHub\\Python\\Gerando Certificados Em Massa Com Python\\DadosAlunosEmail.xlsx"
planilhadadosaluno = load_workbook(arquivoalunos)

sheet_selecionada = planilhadadosaluno["Nomes"]

for linha in range(2, len(sheet_selecionada["A"]) + 1):

    arquivoword = Document("C:\\Users\\georg\\Documents\\Python - GitHub\\Python\\Gerando Certificados Em Massa Com Python\\Certificado3.docx")

    estilo = arquivoword.styles["Normal"]

    #pegando o nome do aluno qnd passar na celula
    nomealuno = sheet_selecionada['A%s' % linha].value
    dia = sheet_selecionada['B%s' % linha].value
    mes = sheet_selecionada['C%s' % linha].value
    ano = sheet_selecionada['D%s' % linha].value
    nomecurso = sheet_selecionada['E%s' % linha].value
    nomeinstrutor = sheet_selecionada['F%s' % linha].value
    emailaluno = sheet_selecionada['G%s' % linha].value

    for paragrafo in arquivoword.paragraphs:

        if "@nome" in paragrafo.text:
            paragrafo.text = nomealuno
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

        paragrafop1 = "Concluiu com sucesso o curso de"
        paragrafop2 = ", com a carga horária de 20 horas, promovido pela escola de Cursos Online em "
        paragrafocomplemento = f"{paragrafop2} {dia} de {mes} de {ano}"

        if "escola" in paragrafo.text:
            paragrafo.text = paragrafop1
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)
            adicionanovapalavra = paragrafo.add_run(nomecurso)
            adicionanovapalavra.font.color.rgb = RGBColor(255,122,157) #cor
            adicionanovapalavra.underline = True #sublinhado
            adicionanovapalavra.bold = True #negrito
            adicionanovapalavra = paragrafo.add_run(paragrafocomplemento)
            adicionanovapalavra.font.color.rgb = RGBColor(0, 0, 0)

        if "Instrutor" in paragrafo.text:
            paragrafo.text = nomeinstrutor + " - Instrutor"
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

    caminhocertificados = "C:\\Users\\georg\\Downloads\\" + nomealuno + ".docx"
    arquivoword.save(caminhocertificados)

    emailoutlook = outlook.CreateItem(0)
    emailoutlook.To = emailaluno
    emailoutlook.Subject = "Certificado " + nomealuno
    emailoutlook.HTMLBody = f"""
        <p> Boa noite {nomealuno}!</p>
        <p> Segue seu <b>certificado:</b></p>
        <p> Atenciosamente, </p>
        <p> <img src  = "C:\\Users\\georg\\Documents\\Python - GitHub\\Python\\Gerando Certificados Em Massa Com Python\\imgEmail.jpg"></p>
    """

    #adiciona anexo no email
    emailoutlook.Attachments.Add(caminhocertificados)

    #ou .send() /.send
    emailoutlook.save()

print("Certificados gerados com sucesso!")